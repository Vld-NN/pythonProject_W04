import tkinter as tk
from tkinter import ttk
import docx


def save_data():
    name = entry_name.get()
    age = entry_age.get()
    city = entry_city.get()
    email = entry_email.get()
    phone = entry_phone.get()

    print(f'Имя: {name}, Возраст: {age}, Город: {city}, Email: {email}, Телефон: {phone}')

    # Генерируем документ Word
    generate_word_document(name, age, city, email, phone)


def generate_word_document(name, age, city, email, phone):
    document = docx.Document()  # Создаем новый документ

    # Добавляем заголовок
    document.add_heading('Анкета участника', level=0)

    # Добавляем параграфы с информацией
    document.add_paragraph(f'Имя: {name}')
    document.add_paragraph(f'Возраст: {age}')
    document.add_paragraph(f'Город: {city}')
    document.add_paragraph(f'Email: {email}')
    document.add_paragraph(f'Телефон: {phone}')

    # Сохраняем документ
    document.save('анкета.docx')

    print("Документ успешно создан!")


# Создаём главное окно приложения
root = tk.Tk()
root.title('Заполнение анкеты')

# Поле для имени
label_name = ttk.Label(root, text='Введите ваше имя:')
label_name.pack(padx=10, pady=10)
entry_name = ttk.Entry(root)
entry_name.pack(padx=10, pady=10)

# Поле для возраста
label_age = ttk.Label(root, text='Введите ваш возраст:')
label_age.pack(padx=10, pady=10)
entry_age = ttk.Entry(root)
entry_age.pack(padx=10, pady=10)

# Поле для города
label_city = ttk.Label(root, text='Введите ваш город:')
label_city.pack(padx=10, pady=10)
entry_city = ttk.Entry(root)
entry_city.pack(padx=10, pady=10)

# Поле для адреса электронной почты
label_email = ttk.Label(root, text='Введите ваш email:')
label_email.pack(padx=10, pady=10)
entry_email = ttk.Entry(root)
entry_email.pack(padx=10, pady=10)

# Поле для номера телефона
label_phone = ttk.Label(root, text='Введите ваш телефон:')
label_phone.pack(padx=10, pady=10)
entry_phone = ttk.Entry(root)
entry_phone.pack(padx=10, pady=10)

# Кнопка сохранения данных
button_save = ttk.Button(root, text='Сохранить', command=save_data)
button_save.pack(padx=10, pady=10)

# Запуск главного цикла приложения
root.mainloop()
